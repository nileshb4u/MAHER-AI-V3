"""
Word/Docx Utilities Tool for MAHER Orchestrator

Provides comprehensive Word document manipulation capabilities including:
- Create Word documents
- Add/modify tables
- Extract text and metadata
- Extract headings and structure
- Convert to PDF
- Track changes
- Add formatting and styles
"""

from typing import Dict, Any, List, Optional
import os
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    import mammoth
    DEPENDENCIES_AVAILABLE = True
except ImportError as e:
    DEPENDENCIES_AVAILABLE = False
    IMPORT_ERROR = str(e)


def create_word_document(
    output_file: str,
    title: str = "",
    content: Optional[List[Dict[str, Any]]] = None
) -> Dict[str, Any]:
    """
    Create a new Word document with specified content

    Args:
        output_file: Path for output Word document
        title: Document title
        content: List of content items with type and data
            Example: [
                {"type": "heading", "text": "Title", "level": 1},
                {"type": "paragraph", "text": "Content"},
                {"type": "table", "data": [[row1], [row2]]}
            ]

    Returns:
        Result dictionary with success status
    """
    if not DEPENDENCIES_AVAILABLE:
        return {
            "success": False,
            "error": f"Required dependencies not available: {IMPORT_ERROR}"
        }

    try:
        doc = Document()

        # Add title if provided
        if title:
            doc.add_heading(title, 0)

        # Add content items
        if content:
            for item in content:
                item_type = item.get("type", "paragraph")

                if item_type == "heading":
                    level = item.get("level", 1)
                    doc.add_heading(item.get("text", ""), level=level)

                elif item_type == "paragraph":
                    p = doc.add_paragraph(item.get("text", ""))

                    # Apply formatting if specified
                    if item.get("bold"):
                        p.runs[0].bold = True
                    if item.get("italic"):
                        p.runs[0].italic = True
                    if item.get("alignment"):
                        alignment = item["alignment"]
                        if alignment == "center":
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif alignment == "right":
                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                elif item_type == "table":
                    table_data = item.get("data", [])
                    if table_data:
                        rows = len(table_data)
                        cols = len(table_data[0]) if table_data else 0
                        table = doc.add_table(rows=rows, cols=cols)
                        table.style = 'Light Grid Accent 1'

                        for i, row_data in enumerate(table_data):
                            for j, cell_data in enumerate(row_data):
                                table.rows[i].cells[j].text = str(cell_data)

                elif item_type == "bullet_list":
                    items = item.get("items", [])
                    for list_item in items:
                        doc.add_paragraph(list_item, style='List Bullet')

                elif item_type == "numbered_list":
                    items = item.get("items", [])
                    for list_item in items:
                        doc.add_paragraph(list_item, style='List Number')

                elif item_type == "page_break":
                    doc.add_page_break()

        # Save document
        doc.save(output_file)

        return {
            "success": True,
            "output_file": output_file,
            "message": "Word document created successfully"
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Document creation failed: {str(e)}"
        }


def add_table_to_word(
    input_file: str,
    output_file: str,
    table_data: List[List[Any]],
    position: str = "end",
    style: str = "Light Grid Accent 1"
) -> Dict[str, Any]:
    """
    Add a table to an existing Word document

    Args:
        input_file: Path to input Word document
        output_file: Path for output document
        table_data: 2D list of table data
        position: Where to insert table - 'start', 'end', or paragraph index
        style: Table style name

    Returns:
        Result dictionary
    """
    if not DEPENDENCIES_AVAILABLE:
        return {
            "success": False,
            "error": f"Required dependencies not available: {IMPORT_ERROR}"
        }

    try:
        doc = Document(input_file)

        # Create table
        rows = len(table_data)
        cols = len(table_data[0]) if table_data else 0
        table = doc.add_table(rows=rows, cols=cols)
        table.style = style

        # Populate table
        for i, row_data in enumerate(table_data):
            for j, cell_data in enumerate(row_data):
                table.rows[i].cells[j].text = str(cell_data)

        # Save document
        doc.save(output_file)

        return {
            "success": True,
            "output_file": output_file,
            "rows": rows,
            "columns": cols,
            "message": f"Table added with {rows} rows and {cols} columns"
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Failed to add table: {str(e)}"
        }


def extract_text_from_word(
    input_file: str,
    include_formatting: bool = False
) -> Dict[str, Any]:
    """
    Extract text from Word document

    Args:
        input_file: Path to Word document
        include_formatting: Whether to include formatting information

    Returns:
        Extracted text and metadata
    """
    if not DEPENDENCIES_AVAILABLE:
        return {
            "success": False,
            "error": f"Required dependencies not available: {IMPORT_ERROR}"
        }

    try:
        doc = Document(input_file)

        paragraphs = []
        for para in doc.paragraphs:
            para_data = {
                "text": para.text,
            }

            if include_formatting:
                para_data["style"] = para.style.name
                para_data["alignment"] = str(para.alignment)

            paragraphs.append(para_data)

        # Extract full text
        full_text = "\n".join([p.text for p in doc.paragraphs])

        return {
            "success": True,
            "text": full_text,
            "paragraphs": paragraphs,
            "paragraph_count": len(paragraphs),
            "message": "Text extracted successfully"
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Text extraction failed: {str(e)}"
        }


def extract_headings_from_word(
    input_file: str
) -> Dict[str, Any]:
    """
    Extract document structure (headings) from Word document

    Args:
        input_file: Path to Word document

    Returns:
        Document structure with headings
    """
    if not DEPENDENCIES_AVAILABLE:
        return {
            "success": False,
            "error": f"Required dependencies not available: {IMPORT_ERROR}"
        }

    try:
        doc = Document(input_file)

        headings = []
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 0
                headings.append({
                    "text": para.text,
                    "level": level,
                    "style": para.style.name
                })

        # Create document outline
        outline = []
        for heading in headings:
            outline.append("  " * (heading["level"] - 1) + heading["text"])

        return {
            "success": True,
            "headings": headings,
            "outline": "\n".join(outline),
            "total_headings": len(headings),
            "message": "Document structure extracted"
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Heading extraction failed: {str(e)}"
        }


def extract_tables_from_word(
    input_file: str
) -> Dict[str, Any]:
    """
    Extract all tables from Word document

    Args:
        input_file: Path to Word document

    Returns:
        All tables data
    """
    if not DEPENDENCIES_AVAILABLE:
        return {
            "success": False,
            "error": f"Required dependencies not available: {IMPORT_ERROR}"
        }

    try:
        doc = Document(input_file)

        tables_data = []
        for table_idx, table in enumerate(doc.tables):
            table_content = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_content.append(row_data)

            tables_data.append({
                "table_index": table_idx + 1,
                "data": table_content,
                "rows": len(table_content),
                "columns": len(table_content[0]) if table_content else 0
            })

        return {
            "success": True,
            "tables": tables_data,
            "total_tables": len(tables_data),
            "message": f"Extracted {len(tables_data)} tables"
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Table extraction failed: {str(e)}"
        }


def word_to_pdf(
    input_file: str,
    output_file: str,
    method: str = "direct"
) -> Dict[str, Any]:
    """
    Convert Word document to PDF

    Args:
        input_file: Path to Word document
        output_file: Path for output PDF
        method: Conversion method - 'direct' (no Office needed) or 'docx2pdf' (requires Office)

    Returns:
        Result dictionary
    """
    if not DEPENDENCIES_AVAILABLE:
        return {
            "success": False,
            "error": f"Required dependencies not available: {IMPORT_ERROR}"
        }

    try:
        if method == "direct":
            # Use direct conversion (no MS Office needed)
            from tools.word_to_pdf_direct import word_to_pdf_direct
            return word_to_pdf_direct(input_file, output_file)

        elif method == "docx2pdf":
            # Try docx2pdf method (requires MS Office/LibreOffice)
            try:
                from docx2pdf import convert
                convert(input_file, output_file)

                return {
                    "success": True,
                    "output_file": output_file,
                    "method": "docx2pdf",
                    "message": "Successfully converted Word to PDF using docx2pdf"
                }

            except ImportError:
                # Fallback to direct method
                from tools.word_to_pdf_direct import word_to_pdf_direct
                return word_to_pdf_direct(input_file, output_file)

        else:
            return {
                "success": False,
                "error": f"Unknown conversion method: {method}",
                "available_methods": ["direct", "docx2pdf"]
            }

    except Exception as e:
        return {
            "success": False,
            "error": f"Conversion failed: {str(e)}"
        }


def word_to_html(
    input_file: str,
    output_file: str
) -> Dict[str, Any]:
    """
    Convert Word document to HTML

    Args:
        input_file: Path to Word document
        output_file: Path for output HTML

    Returns:
        Result dictionary
    """
    if not DEPENDENCIES_AVAILABLE:
        return {
            "success": False,
            "error": f"Required dependencies not available: {IMPORT_ERROR}"
        }

    try:
        with open(input_file, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value
            messages = result.messages

        with open(output_file, "w", encoding="utf-8") as html_file:
            html_file.write(html)

        return {
            "success": True,
            "output_file": output_file,
            "warnings": len(messages),
            "message": "Successfully converted Word to HTML"
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Conversion failed: {str(e)}"
        }


def modify_word_document(
    input_file: str,
    output_file: str,
    modifications: List[Dict[str, Any]]
) -> Dict[str, Any]:
    """
    Modify existing Word document

    Args:
        input_file: Path to input Word document
        output_file: Path for output document
        modifications: List of modifications to apply
            Example: [
                {"action": "add_paragraph", "text": "New content"},
                {"action": "add_heading", "text": "New Section", "level": 2},
                {"action": "replace_text", "find": "old", "replace": "new"}
            ]

    Returns:
        Result dictionary
    """
    if not DEPENDENCIES_AVAILABLE:
        return {
            "success": False,
            "error": f"Required dependencies not available: {IMPORT_ERROR}"
        }

    try:
        doc = Document(input_file)

        for mod in modifications:
            action = mod.get("action", "")

            if action == "add_paragraph":
                doc.add_paragraph(mod.get("text", ""))

            elif action == "add_heading":
                level = mod.get("level", 1)
                doc.add_heading(mod.get("text", ""), level=level)

            elif action == "replace_text":
                find_text = mod.get("find", "")
                replace_text = mod.get("replace", "")

                for para in doc.paragraphs:
                    if find_text in para.text:
                        para.text = para.text.replace(find_text, replace_text)

            elif action == "add_page_break":
                doc.add_page_break()

        doc.save(output_file)

        return {
            "success": True,
            "output_file": output_file,
            "modifications_applied": len(modifications),
            "message": "Document modified successfully"
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Modification failed: {str(e)}"
        }


def get_word_document_info(
    input_file: str
) -> Dict[str, Any]:
    """
    Get metadata and statistics about Word document

    Args:
        input_file: Path to Word document

    Returns:
        Document metadata and statistics
    """
    if not DEPENDENCIES_AVAILABLE:
        return {
            "success": False,
            "error": f"Required dependencies not available: {IMPORT_ERROR}"
        }

    try:
        doc = Document(input_file)
        core_props = doc.core_properties

        return {
            "success": True,
            "file_path": input_file,
            "metadata": {
                "title": core_props.title or "N/A",
                "author": core_props.author or "N/A",
                "subject": core_props.subject or "N/A",
                "keywords": core_props.keywords or "N/A",
                "created": str(core_props.created) if core_props.created else "N/A",
                "modified": str(core_props.modified) if core_props.modified else "N/A",
                "last_modified_by": core_props.last_modified_by or "N/A",
            },
            "statistics": {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections),
                "file_size": os.path.getsize(input_file)
            }
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Failed to get document info: {str(e)}"
        }
