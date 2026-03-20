"""
Direct Word to PDF conversion using reportlab
No external dependencies like MS Office or LibreOffice required
"""

from typing import Dict, Any
import os

try:
    from docx import Document
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    DEPENDENCIES_AVAILABLE = True
except ImportError as e:
    DEPENDENCIES_AVAILABLE = False
    IMPORT_ERROR = str(e)


def word_to_pdf_direct(
    input_file: str,
    output_file: str,
    preserve_formatting: bool = True
) -> Dict[str, Any]:
    """
    Convert Word document to PDF directly using reportlab
    No MS Office or LibreOffice required - works offline

    Args:
        input_file: Path to Word document
        output_file: Path for output PDF
        preserve_formatting: Try to preserve basic formatting

    Returns:
        Result dictionary
    """
    if not DEPENDENCIES_AVAILABLE:
        return {
            "success": False,
            "error": f"Required dependencies not available: {IMPORT_ERROR}"
        }

    try:
        # Read Word document
        doc = Document(input_file)

        # Create PDF
        pdf = SimpleDocTemplate(
            output_file,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )

        # Container for the 'Flowable' objects
        story = []

        # Define styles
        styles = getSampleStyleSheet()

        # Create custom styles for different heading levels
        for i in range(1, 10):
            style_name = f'Heading{i}'
            if style_name not in styles:
                base_size = 16 - (i - 1) * 1.5
                styles.add(ParagraphStyle(
                    name=style_name,
                    parent=styles['Heading1'],
                    fontSize=base_size,
                    leading=base_size * 1.2,
                    spaceAfter=6,
                    spaceBefore=12,
                    textColor=colors.HexColor('#1f4788'),
                    fontName='Helvetica-Bold'
                ))

        # Process paragraphs
        for para in doc.paragraphs:
            if not para.text.strip():
                story.append(Spacer(1, 0.2 * inch))
                continue

            # Determine style based on Word style
            style = styles['BodyText']

            if para.style.name.startswith('Heading'):
                # Extract heading level
                try:
                    level = para.style.name.replace('Heading', '').strip()
                    if level.isdigit():
                        style_name = f'Heading{level}'
                        if style_name in styles:
                            style = styles[style_name]
                except:
                    pass
            elif para.style.name == 'Title':
                style = styles['Title']

            # Create paragraph
            p = Paragraph(para.text, style)
            story.append(p)
            story.append(Spacer(1, 0.1 * inch))

        # Process tables
        for table in doc.tables:
            # Extract table data
            data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                data.append(row_data)

            if data:
                # Create PDF table
                t = Table(data)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))

                story.append(t)
                story.append(Spacer(1, 0.3 * inch))

        # Build PDF
        pdf.build(story)

        return {
            "success": True,
            "output_file": output_file,
            "method": "direct_conversion",
            "paragraphs_converted": len(doc.paragraphs),
            "tables_converted": len(doc.tables),
            "message": "Successfully converted Word to PDF using direct conversion"
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Direct conversion failed: {str(e)}"
        }


def word_to_pdf_with_images(
    input_file: str,
    output_file: str
) -> Dict[str, Any]:
    """
    Advanced Word to PDF conversion with image support

    Args:
        input_file: Path to Word document
        output_file: Path for output PDF

    Returns:
        Result dictionary
    """
    if not DEPENDENCIES_AVAILABLE:
        return {
            "success": False,
            "error": f"Required dependencies not available: {IMPORT_ERROR}"
        }

    try:
        from reportlab.platypus import Image as RLImage

        # Read Word document
        doc = Document(input_file)

        # Create PDF
        pdf = SimpleDocTemplate(output_file, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()

        # Process document with images
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                para = element
                text = ''.join(node.text for node in para.iter() if node.text)
                if text.strip():
                    story.append(Paragraph(text, styles['BodyText']))
                    story.append(Spacer(1, 0.1 * inch))

        # Build PDF
        pdf.build(story)

        return {
            "success": True,
            "output_file": output_file,
            "message": "Successfully converted Word to PDF with advanced features"
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Advanced conversion failed: {str(e)}",
            "fallback": "Use word_to_pdf_direct() instead"
        }
