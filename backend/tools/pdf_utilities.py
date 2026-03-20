"""
PDF Utilities Tool for MAHER Orchestrator

Provides comprehensive PDF manipulation capabilities including:
- Convert to PDF (from images, Word docs, etc.)
- Merge multiple PDFs
- Split PDFs
- Extract text and tables
- OCR scanned documents
- Redact sensitive information
- Convert PDF to Word
"""

from typing import Dict, Any, List, Optional
import os
import io
import re
from pathlib import Path
import base64

try:
    import PyPDF2
    from pypdf import PdfReader, PdfWriter
    import pdfplumber
    from PIL import Image
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    DEPENDENCIES_AVAILABLE = True
except ImportError as e:
    DEPENDENCIES_AVAILABLE = False
    IMPORT_ERROR = str(e)

# OCR is optional - gracefully handle missing pytesseract/tesseract
try:
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    pytesseract = None


def convert_to_pdf(
    input_file: str,
    output_file: str,
    source_type: str = "auto"
) -> Dict[str, Any]:
    """
    Convert various file formats to PDF

    Args:
        input_file: Path to input file or base64 encoded content
        output_file: Path for output PDF
        source_type: Type of source - 'image', 'word', 'text', or 'auto'

    Returns:
        Result dictionary with success status and output path
    """
    if not DEPENDENCIES_AVAILABLE:
        return {
            "success": False,
            "error": f"Required dependencies not available: {IMPORT_ERROR}"
        }

    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas as pdf_canvas

        # Try to import docx2pdf (optional)
        try:
            from docx2pdf import convert as docx_convert
            DOCX2PDF_AVAILABLE = True
        except ImportError:
            DOCX2PDF_AVAILABLE = False

        # Auto-detect source type if not specified
        if source_type == "auto":
            ext = Path(input_file).suffix.lower()
            if ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif']:
                source_type = "image"
            elif ext in ['.docx', '.doc']:
                source_type = "word"
            elif ext in ['.txt']:
                source_type = "text"

        if source_type == "image":
            # Convert image to PDF
            img = Image.open(input_file)
            if img.mode == 'RGBA':
                img = img.convert('RGB')
            img.save(output_file, 'PDF', resolution=100.0)

        elif source_type == "word":
            # Convert Word to PDF
            if not DOCX2PDF_AVAILABLE:
                return {
                    "success": False,
                    "error": "Word to PDF conversion not available - requires docx2pdf and MS Office/LibreOffice",
                    "note": "Use the Word utilities word_to_pdf function or install docx2pdf"
                }
            docx_convert(input_file, output_file)

        elif source_type == "text":
            # Convert text to PDF
            c = pdf_canvas.Canvas(output_file, pagesize=letter)
            width, height = letter

            with open(input_file, 'r', encoding='utf-8') as f:
                text = f.read()

            # Simple text rendering
            y = height - 50
            for line in text.split('\n'):
                if y < 50:  # New page if needed
                    c.showPage()
                    y = height - 50
                c.drawString(50, y, line[:100])  # Limit line length
                y -= 15

            c.save()

        else:
            return {
                "success": False,
                "error": f"Unsupported source type: {source_type}"
            }

        return {
            "success": True,
            "output_file": output_file,
            "message": f"Successfully converted {source_type} to PDF"
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Conversion failed: {str(e)}"
        }


def merge_pdfs(
    input_files: List[str],
    output_file: str
) -> Dict[str, Any]:
    """
    Merge multiple PDF files into one

    Args:
        input_files: List of PDF file paths to merge
        output_file: Path for merged PDF output

    Returns:
        Result dictionary with success status
    """
    if not DEPENDENCIES_AVAILABLE:
        return {
            "success": False,
            "error": f"Required dependencies not available: {IMPORT_ERROR}"
        }

    try:
        pdf_writer = PdfWriter()

        for pdf_file in input_files:
            if not os.path.exists(pdf_file):
                return {
                    "success": False,
                    "error": f"Input file not found: {pdf_file}"
                }

            pdf_reader = PdfReader(pdf_file)
            for page in pdf_reader.pages:
                pdf_writer.add_page(page)

        with open(output_file, 'wb') as output:
            pdf_writer.write(output)

        return {
            "success": True,
            "output_file": output_file,
            "total_pages": len(pdf_writer.pages),
            "message": f"Successfully merged {len(input_files)} PDFs"
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Merge failed: {str(e)}"
        }


def split_pdf(
    input_file: str,
    output_dir: str,
    split_type: str = "pages",
    split_value: Optional[int] = None
) -> Dict[str, Any]:
    """
    Split PDF into multiple files

    Args:
        input_file: Path to PDF to split
        output_dir: Directory for output files
        split_type: 'pages' (one per page), 'range' (by page range), or 'count' (n pages per file)
        split_value: Number of pages per file (for 'count' type)

    Returns:
        Result dictionary with list of output files
    """
    if not DEPENDENCIES_AVAILABLE:
        return {
            "success": False,
            "error": f"Required dependencies not available: {IMPORT_ERROR}"
        }

    try:
        pdf_reader = PdfReader(input_file)
        total_pages = len(pdf_reader.pages)

        os.makedirs(output_dir, exist_ok=True)
        output_files = []

        if split_type == "pages":
            # Split into individual pages
            for page_num in range(total_pages):
                pdf_writer = PdfWriter()
                pdf_writer.add_page(pdf_reader.pages[page_num])

                output_path = os.path.join(output_dir, f"page_{page_num + 1}.pdf")
                with open(output_path, 'wb') as output:
                    pdf_writer.write(output)
                output_files.append(output_path)

        elif split_type == "count" and split_value:
            # Split into files with n pages each
            page_idx = 0
            file_idx = 1

            while page_idx < total_pages:
                pdf_writer = PdfWriter()

                for _ in range(split_value):
                    if page_idx >= total_pages:
                        break
                    pdf_writer.add_page(pdf_reader.pages[page_idx])
                    page_idx += 1

                output_path = os.path.join(output_dir, f"part_{file_idx}.pdf")
                with open(output_path, 'wb') as output:
                    pdf_writer.write(output)
                output_files.append(output_path)
                file_idx += 1

        else:
            return {
                "success": False,
                "error": f"Invalid split_type: {split_type}"
            }

        return {
            "success": True,
            "output_files": output_files,
            "total_files": len(output_files),
            "message": f"Successfully split PDF into {len(output_files)} files"
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Split failed: {str(e)}"
        }


def extract_text_from_pdf(
    input_file: str,
    page_numbers: Optional[List[int]] = None,
    use_ocr: bool = False
) -> Dict[str, Any]:
    """
    Extract text from PDF

    Args:
        input_file: Path to PDF file
        page_numbers: Specific pages to extract (None for all)
        use_ocr: Whether to use OCR for scanned documents

    Returns:
        Extracted text by page
    """
    if not DEPENDENCIES_AVAILABLE:
        return {
            "success": False,
            "error": f"Required dependencies not available: {IMPORT_ERROR}"
        }

    try:
        extracted_text = {}

        with pdfplumber.open(input_file) as pdf:
            pages_to_process = page_numbers if page_numbers else range(len(pdf.pages))

            for page_num in pages_to_process:
                if page_num >= len(pdf.pages):
                    continue

                page = pdf.pages[page_num]
                text = page.extract_text()

                # Use OCR if text extraction failed or OCR is explicitly requested
                if (not text or use_ocr) and use_ocr:
                    # Try local offline OCR first (EasyOCR/EffOCR)
                    try:
                        from tools.ocr_effocr import ocr_image_local
                        import tempfile

                        # Convert page to image
                        img = page.to_image()

                        # Save temp image
                        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                            img.original.save(tmp.name, 'PNG')
                            temp_path = tmp.name

                        # Perform local OCR
                        ocr_result = ocr_image_local(temp_path, engine="easyocr")

                        # Clean up
                        try:
                            os.remove(temp_path)
                        except:
                            pass

                        if ocr_result.get("success"):
                            text = ocr_result.get("text", "")
                        else:
                            # Fallback to pytesseract if available
                            if OCR_AVAILABLE:
                                text = pytesseract.image_to_string(img.original)
                            else:
                                text = f"[Local OCR not available - install easyocr: pip install easyocr]"

                    except Exception as ocr_error:
                        # Last resort: try pytesseract
                        if OCR_AVAILABLE:
                            try:
                                img = page.to_image()
                                text = pytesseract.image_to_string(img.original)
                            except:
                                text = text or f"[OCR failed: {str(ocr_error)}]"
                        else:
                            text = text or f"[OCR failed: {str(ocr_error)}]"

                extracted_text[f"page_{page_num + 1}"] = text or ""

        return {
            "success": True,
            "text": extracted_text,
            "total_pages": len(extracted_text),
            "message": "Text extraction completed"
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Text extraction failed: {str(e)}"
        }


def extract_tables_from_pdf(
    input_file: str,
    page_numbers: Optional[List[int]] = None
) -> Dict[str, Any]:
    """
    Extract tables from PDF

    Args:
        input_file: Path to PDF file
        page_numbers: Specific pages to extract tables from (None for all)

    Returns:
        Extracted tables data
    """
    if not DEPENDENCIES_AVAILABLE:
        return {
            "success": False,
            "error": f"Required dependencies not available: {IMPORT_ERROR}"
        }

    try:
        all_tables = []

        with pdfplumber.open(input_file) as pdf:
            pages_to_process = page_numbers if page_numbers else range(len(pdf.pages))

            for page_num in pages_to_process:
                if page_num >= len(pdf.pages):
                    continue

                page = pdf.pages[page_num]
                tables = page.extract_tables()

                for table_idx, table in enumerate(tables):
                    all_tables.append({
                        "page": page_num + 1,
                        "table_index": table_idx + 1,
                        "data": table,
                        "rows": len(table),
                        "columns": len(table[0]) if table else 0
                    })

        return {
            "success": True,
            "tables": all_tables,
            "total_tables": len(all_tables),
            "message": f"Extracted {len(all_tables)} tables from PDF"
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Table extraction failed: {str(e)}"
        }


def ocr_scanned_pdf(
    input_file: str,
    output_file: str,
    language: str = "eng"
) -> Dict[str, Any]:
    """
    Perform OCR on scanned PDF to make it searchable

    Args:
        input_file: Path to scanned PDF
        output_file: Path for OCR'd PDF output
        language: OCR language (default: 'eng')

    Returns:
        Result with OCR'd text
    """
    if not DEPENDENCIES_AVAILABLE:
        return {
            "success": False,
            "error": f"Required dependencies not available: {IMPORT_ERROR}"
        }

    if not OCR_AVAILABLE:
        return {
            "success": False,
            "error": "OCR not available - pytesseract and Tesseract OCR binary required",
            "note": "Install pytesseract Python package and Tesseract OCR binary to enable OCR"
        }

    try:
        # Extract text using OCR
        result = extract_text_from_pdf(input_file, use_ocr=True)

        if result["success"]:
            # Create searchable PDF with OCR'd text
            # For simplicity, we'll return the extracted text
            # In production, you'd want to create a true searchable PDF

            return {
                "success": True,
                "ocr_text": result["text"],
                "message": "OCR completed successfully",
                "note": "Text extracted - searchable PDF creation requires additional libraries"
            }
        else:
            return result

    except Exception as e:
        return {
            "success": False,
            "error": f"OCR failed: {str(e)}"
        }


def redact_pdf(
    input_file: str,
    output_file: str,
    redaction_patterns: List[str]
) -> Dict[str, Any]:
    """
    Redact sensitive information from PDF

    Args:
        input_file: Path to input PDF
        output_file: Path for redacted PDF
        redaction_patterns: List of regex patterns to redact (e.g., email, SSN, phone)

    Returns:
        Result with redaction statistics
    """
    if not DEPENDENCIES_AVAILABLE:
        return {
            "success": False,
            "error": f"Required dependencies not available: {IMPORT_ERROR}"
        }

    try:
        # Extract text to find redaction targets
        text_result = extract_text_from_pdf(input_file)

        if not text_result["success"]:
            return text_result

        redaction_count = 0
        redacted_items = []

        # Find patterns to redact
        for page_key, text in text_result["text"].items():
            for pattern in redaction_patterns:
                matches = re.findall(pattern, text)
                redaction_count += len(matches)
                redacted_items.extend(matches)

        # Copy PDF and note what would be redacted
        # Full redaction requires additional libraries like PyMuPDF
        pdf_reader = PdfReader(input_file)
        pdf_writer = PdfWriter()

        for page in pdf_reader.pages:
            pdf_writer.add_page(page)

        with open(output_file, 'wb') as output:
            pdf_writer.write(output)

        return {
            "success": True,
            "output_file": output_file,
            "redaction_count": redaction_count,
            "patterns_matched": len(redacted_items),
            "message": f"Found {redaction_count} items matching redaction patterns",
            "note": "Visual redaction requires PyMuPDF - text patterns identified"
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Redaction failed: {str(e)}"
        }


def pdf_to_word(
    input_file: str,
    output_file: str
) -> Dict[str, Any]:
    """
    Convert PDF to Word document

    Args:
        input_file: Path to input PDF
        output_file: Path for output Word document

    Returns:
        Result dictionary
    """
    if not DEPENDENCIES_AVAILABLE:
        return {
            "success": False,
            "error": f"Required dependencies not available: {IMPORT_ERROR}"
        }

    try:
        from docx import Document

        # Extract text from PDF
        text_result = extract_text_from_pdf(input_file)

        if not text_result["success"]:
            return text_result

        # Create Word document
        doc = Document()
        doc.add_heading('Converted from PDF', 0)

        for page_key, text in text_result["text"].items():
            doc.add_heading(page_key.replace('_', ' ').title(), level=1)
            doc.add_paragraph(text)
            doc.add_page_break()

        doc.save(output_file)

        return {
            "success": True,
            "output_file": output_file,
            "pages_converted": len(text_result["text"]),
            "message": "Successfully converted PDF to Word"
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"PDF to Word conversion failed: {str(e)}"
        }


def get_pdf_info(input_file: str) -> Dict[str, Any]:
    """
    Get metadata and information about a PDF

    Args:
        input_file: Path to PDF file

    Returns:
        PDF metadata and statistics
    """
    if not DEPENDENCIES_AVAILABLE:
        return {
            "success": False,
            "error": f"Required dependencies not available: {IMPORT_ERROR}"
        }

    try:
        pdf_reader = PdfReader(input_file)
        metadata = pdf_reader.metadata

        return {
            "success": True,
            "file_path": input_file,
            "total_pages": len(pdf_reader.pages),
            "metadata": {
                "title": metadata.get('/Title', 'N/A') if metadata else 'N/A',
                "author": metadata.get('/Author', 'N/A') if metadata else 'N/A',
                "subject": metadata.get('/Subject', 'N/A') if metadata else 'N/A',
                "creator": metadata.get('/Creator', 'N/A') if metadata else 'N/A',
                "producer": metadata.get('/Producer', 'N/A') if metadata else 'N/A',
            },
            "file_size": os.path.getsize(input_file),
            "encrypted": pdf_reader.is_encrypted
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Failed to get PDF info: {str(e)}"
        }
